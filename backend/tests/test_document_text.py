"""Orion reads the formats the job produces, not only the ones a person types.

The GPS export is an Excel workbook, the medical report is a PDF, the scouting
summary is a Word file. Before this, all three had to be converted by hand
before Orion could see them - the tool asking its user to do the tool's work.

Every file in these tests is built for real by the same libraries that will
read it: a fixture string asserting "it parses" would only prove the fixture
matches the parser. What is checked is that content put into a document comes
back out of it, and that an unreadable file says why in words the person
holding it can act on.
"""

from __future__ import annotations

import io
import unittest

from backend.app.services.document_text import (
    MAX_EXTRACTED_CHARACTERS,
    MAX_SHEET_ROWS,
    SUPPORTED_EXTENSIONS,
    DocumentExtractionError,
    extract_text,
)


def _pdf_bytes(pages: list[str]) -> bytes:
    """A real PDF, assembled by hand: catalog, pages, fonts, xref and all.

    pypdf writes documents but does not typeset, so a fixture built with it
    carries no extractable text - which would make the extractor look broken
    when it is not, or worse, pass while reading nothing. Writing the file the
    way a PDF actually is means the test exercises the same path a report from
    a clinic does.
    """

    objects: list[bytes] = []
    page_count = len(pages)
    # Object numbering: 1 catalog, 2 pages, 3 font, then a page and a content
    # stream per text.
    page_ids = [4 + index * 2 for index in range(page_count)]
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects.append(
        f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode("latin-1")
    )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for index, text in enumerate(pages):
        content_id = page_ids[index] + 1
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] "
                f"/Contents {content_id} 0 R "
                f"/Resources << /Font << /F1 3 0 R >> >> >>"
            ).encode("latin-1")
        )
        stream = f"BT /F1 12 Tf 20 250 Td ({text}) Tj ET".encode("latin-1")
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode("latin-1")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode("latin-1") + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("latin-1")
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode("latin-1")
    return bytes(out)


def _xlsx_bytes(sheets: dict[str, list[list[object]]]) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.remove(workbook.active)
    for title, rows in sheets.items():
        sheet = workbook.create_sheet(title=title)
        for row in rows:
            sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx_bytes(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        created = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for cell_index, value in enumerate(row):
                created.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class PdfTests(unittest.TestCase):
    def test_a_medical_report_becomes_readable_text(self) -> None:
        data = _pdf_bytes(["Fascia plantar sin signos de fascitis", "Alta deportiva"])
        text = extract_text("informe.pdf", data)
        self.assertIn("Fascia plantar", text)
        self.assertIn("Alta deportiva", text)
        # Page markers, so an answer can say where in the report something is.
        self.assertIn("[Página 1]", text)
        self.assertIn("[Página 2]", text)

    def test_a_scanned_pdf_says_so_instead_of_storing_nothing(self) -> None:
        """A scan is a picture of text. Storing it empty answers nothing, silently."""

        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)

        with self.assertRaises(DocumentExtractionError) as caught:
            extract_text("escaneado.pdf", buffer.getvalue())
        self.assertIn("escaneados", str(caught.exception))

    def test_a_damaged_pdf_reports_the_problem(self) -> None:
        with self.assertRaises(DocumentExtractionError) as caught:
            extract_text("roto.pdf", b"%PDF-1.4 esto no es un PDF completo")
        self.assertIn("no se pudo leer", str(caught.exception).lower())


class ExcelTests(unittest.TestCase):
    def test_a_gps_export_arrives_as_csv_on_the_path_orion_computes_over(self) -> None:
        """A workbook becomes CSV so filters and averages already work on it."""

        data = _xlsx_bytes(
            {
                "Sesion": [
                    ["Jugador", "Distancia", "HSR"],
                    ["A", 9800, 420],
                    ["B", 10250, 515],
                ]
            }
        )
        text = extract_text("gps.xlsx", data)
        # No annotation of ours in a single-sheet workbook: it is CSV, so the
        # deterministic tool parses it exactly as it parses an exported file.
        self.assertNotIn("[Hoja:", text)
        self.assertIn("Jugador,Distancia,HSR", text)
        self.assertIn("A,9800,420", text)

    def test_every_sheet_is_read_not_only_the_first(self) -> None:
        data = _xlsx_bytes(
            {
                "Enero": [["Jugador", "Carga"], ["A", 300]],
                "Febrero": [["Jugador", "Carga"], ["A", 415]],
            }
        )
        text = extract_text("cargas.xlsx", data)
        self.assertIn("[Hoja: Enero]", text)
        self.assertIn("[Hoja: Febrero]", text)
        self.assertIn("415", text)

    def test_a_formula_arrives_as_its_result_not_its_source(self) -> None:
        """What the person reading the sheet sees is the number, not "=B2*2"."""

        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Datos"
        sheet.append(["Jugador", "Carga"])
        sheet.append(["A", 300])
        # openpyxl writes no cached value for a formula it never calculated, so
        # data_only yields None - which must come through as an empty cell, not
        # as the formula text leaking into the evidence.
        sheet["C2"] = "=B2*2"
        buffer = io.BytesIO()
        workbook.save(buffer)

        text = extract_text("formula.xlsx", buffer.getvalue())
        self.assertNotIn("=B2*2", text)

    def test_a_huge_sheet_is_bounded_rather_than_read_whole(self) -> None:
        rows = [["Jugador", "Carga"]] + [["A", index] for index in range(MAX_SHEET_ROWS + 500)]
        text = extract_text("enorme.xlsx", _xlsx_bytes({"Datos": rows}))
        self.assertLessEqual(len(text), MAX_EXTRACTED_CHARACTERS + 200)

    def test_an_old_xls_is_told_what_to_do_about_it(self) -> None:
        with self.assertRaises(DocumentExtractionError) as caught:
            extract_text("viejo.xlsx", b"\xd0\xcf\x11\xe0 formato binario antiguo")
        self.assertIn("CSV", str(caught.exception))


class WordTests(unittest.TestCase):
    def test_a_scouting_report_becomes_text(self) -> None:
        data = _docx_bytes(["Informe de scouting", "Buen juego aereo."])
        text = extract_text("scouting.docx", data)
        self.assertIn("Informe de scouting", text)
        self.assertIn("Buen juego aereo.", text)

    def test_the_numbers_in_a_table_are_not_dropped(self) -> None:
        """A report's figures live in its tables, outside `paragraphs`.

        Reading only paragraphs would silently discard exactly the part someone
        uploads the report to ask about.
        """

        data = _docx_bytes(
            ["Resumen de test"],
            table=[["Test", "Resultado"], ["McGill anterior", "94 s"]],
        )
        text = extract_text("test.docx", data)
        self.assertIn("McGill anterior", text)
        self.assertIn("94 s", text)


class FormatRoutingTests(unittest.TestCase):
    def test_plain_formats_still_work_exactly_as_before(self) -> None:
        self.assertIn("Jugador", extract_text("gps.csv", b"Jugador,Carga\nA,300"))
        self.assertIn("hola", extract_text("nota.txt", "hola".encode("utf-8")))

    def test_a_file_saved_from_excel_in_latin1_is_still_readable(self) -> None:
        """Windows exports are routinely not UTF-8; refusing them is refusing the job."""

        text = extract_text("gps.csv", "Jugador,Posición\nA,Extremo".encode("cp1252"))
        self.assertIn("Posición", text)

    def test_an_unsupported_format_names_what_is_accepted(self) -> None:
        with self.assertRaises(DocumentExtractionError) as caught:
            extract_text("video.mp4", b"\x00\x01\x02")
        message = str(caught.exception)
        self.assertIn(".pdf", message)
        self.assertIn(".xlsx", message)

    def test_an_empty_file_is_rejected_before_anything_is_stored(self) -> None:
        with self.assertRaises(DocumentExtractionError):
            extract_text("vacio.csv", b"")

    def test_the_accepted_list_matches_what_the_router_handles(self) -> None:
        """A format in the list that nothing parses would be a promise to no one."""

        for extension in SUPPORTED_EXTENSIONS:
            with self.subTest(extension=extension):
                try:
                    extract_text(f"archivo{extension}", b"contenido de prueba")
                except DocumentExtractionError as exc:
                    # Failing to parse junk is fine; being told the format is
                    # unsupported is not - that means the list overpromises.
                    self.assertNotIn("todavía no puede leer", str(exc))


if __name__ == "__main__":
    unittest.main()


class UploadRouteTests(unittest.TestCase):
    """The whole path a person uses: pick a file, and Orion can answer about it."""

    def setUp(self) -> None:
        import asyncio
        import tempfile
        from pathlib import Path

        from backend.app.api.routes import require_api_key
        from backend.app.core.config import Settings
        from backend.app.main import app
        from backend.app.services.rate_limit import upload_rate_limiter

        self.app = app
        self.require_api_key = require_api_key
        app.dependency_overrides[require_api_key] = lambda: None
        self._temp = tempfile.TemporaryDirectory()
        self.settings = Settings(
            knowledge_path=str(Path(self._temp.name) / "docs.json")
        )
        asyncio.run(upload_rate_limiter.reset())

    def tearDown(self) -> None:
        import asyncio

        from backend.app.services.rate_limit import upload_rate_limiter

        self.app.dependency_overrides.pop(self.require_api_key, None)
        self._temp.cleanup()
        asyncio.run(upload_rate_limiter.reset())

    def _upload(self, name: str, data: bytes):
        from unittest.mock import patch

        from fastapi.testclient import TestClient

        with (
            patch("backend.app.api.routes.get_settings", return_value=self.settings),
            TestClient(self.app, raise_server_exceptions=False) as client,
        ):
            return client.post(
                "/api/v1/knowledge/files", files={"file": (name, data)}
            )

    def test_an_excel_upload_is_stored_as_searchable_text(self) -> None:
        response = self._upload(
            "gps.xlsx",
            _xlsx_bytes({"Sesion": [["Jugador", "HSR"], ["A", 515]]}),
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["name"], "gps.xlsx")
        self.assertGreater(response.json()["characters"], 0)

    def test_a_pdf_upload_survives_the_round_trip_as_bytes(self) -> None:
        """Binary must reach the parser intact, not decoded through UTF-8."""

        response = self._upload("informe.pdf", _pdf_bytes(["Alta deportiva"]))
        self.assertEqual(response.status_code, 200)

    def test_an_unreadable_file_explains_itself_instead_of_a_500(self) -> None:
        response = self._upload("escaneado.pdf", b"%PDF-1.4 roto")
        self.assertEqual(response.status_code, 422)
        detail = response.json()["detail"]
        self.assertEqual(detail["code"], "document_unreadable")
        self.assertIn("PDF", detail["message"])

    def test_an_unsupported_format_is_refused_before_it_is_stored(self) -> None:
        response = self._upload("clip.mp4", b"\x00\x01\x02\x03")
        self.assertEqual(response.status_code, 422)
        self.assertIn(".xlsx", response.json()["detail"]["message"])


class SpreadsheetIsComputableTests(unittest.TestCase):
    """Extracting the text is only half of it; Orion must be able to use it.

    The first version of this feature stored a workbook's text and stopped
    there. Three checks elsewhere still asked "does the name end in .csv?", so
    the catalog showed no columns, the deterministic CSV tool refused the file
    outright, and retrieval chunked a table as if it were prose. Uploading the
    GPS export worked and answering anything about it did not.
    """

    def _document(self, sheets):
        from backend.app.services.knowledge_base import KnowledgeDocument

        return KnowledgeDocument(
            "1", "gps.xlsx", extract_text("gps.xlsx", _xlsx_bytes(sheets))
        )

    def test_one_sheet_becomes_plain_csv_with_no_annotation_of_ours(self) -> None:
        text = self._document({"Sesion": [["Jugador", "HSR"], ["A", 420]]}).content
        self.assertNotIn("[Hoja:", text)
        self.assertTrue(text.startswith("Jugador,HSR"))

    def test_the_planner_sees_the_columns_of_an_uploaded_workbook(self) -> None:
        from backend.app.services.semantic_orchestrator import document_catalog

        catalog = document_catalog(
            [self._document({"Sesion": [["Jugador", "Distancia", "HSR"], ["A", 1, 2]]})]
        )
        self.assertIn("columnas: Jugador, Distancia, HSR", catalog)

    def test_orion_computes_over_an_uploaded_workbook(self) -> None:
        from backend.app.services.semantic_tools import (
            CsvFilter,
            CsvOperationSpec,
            execute_csv_operation,
        )

        document = self._document(
            {
                "Sesion": [
                    ["Jugador", "HSR"],
                    ["Perez", 420],
                    ["Gomez", 515],
                    ["Perez", 480],
                ]
            }
        )
        result = execute_csv_operation(
            [document],
            CsvOperationSpec(
                document_name="gps.xlsx",
                filters=(CsvFilter("Jugador", "Perez"),),
                value_column="HSR",
                aggregation="average",
            ),
        )
        self.assertIsNone(result.error)
        self.assertIn("average(HSR)=450", result.context)

    def test_a_workbook_is_chunked_per_row_like_a_table(self) -> None:
        from backend.app.services.local_retrieval import retrieve_local_chunks

        document = self._document(
            {"Sesion": [["Jugador", "HSR"], ["Perez", 420], ["Gomez", 515]]}
        )
        chunks = retrieve_local_chunks([document], "HSR de Gomez")
        self.assertTrue(any("Fila CSV" in chunk.content for chunk in chunks))

    def test_several_sheets_stay_labelled_rather_than_pretending_to_be_one_table(
        self,
    ) -> None:
        """Two sheets need not share columns, so merging them would invent a table."""

        text = self._document(
            {
                "Enero": [["Jugador", "Carga"], ["A", 300]],
                "Febrero": [["Test", "Resultado"], ["McGill", "94 s"]],
            }
        ).content
        self.assertIn("[Hoja: Enero]", text)
        self.assertIn("[Hoja: Febrero]", text)
