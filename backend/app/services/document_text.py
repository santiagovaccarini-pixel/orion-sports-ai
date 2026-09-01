"""Turn the files a sports professional actually receives into plain text.

Until this existed Orion read `.txt`, `.md`, `.csv` and `.json` - formats a
person types, not the ones a job produces. The GPS export arrives as an Excel
workbook, the medical report as a PDF, the scouting summary as a Word document.
Asking someone to convert each file by hand before Orion can read it is asking
them to do the work the tool exists to do.

Extraction only. Nothing here decides what a document means or whether it is
relevant: it produces text, and the same retrieval and review stages that
already handle a pasted CSV take it from there. A spreadsheet becomes CSV
precisely so it lands on the CSV path Orion already computes over, rather than
a second parallel one that could drift.

Every parser here reads a file chosen by a person but authored by someone else,
so each one is bounded: a limit on pages, rows, sheets and characters, and a
clear error rather than a half-truth when a file cannot be read. The text a
file yields is treated as external content by the reasoning stages, exactly
like a web page.
"""

from __future__ import annotations

import csv
import io
import zipfile

# Bounds chosen so a real working file passes and a runaway one cannot exhaust
# the server: a season's GPS export is thousands of rows, not millions, and a
# medical report is tens of pages, not thousands.
MAX_PDF_PAGES = 300
MAX_SHEET_ROWS = 20_000
MAX_SHEETS = 20
MAX_DOCX_PARAGRAPHS = 20_000
# Matches KnowledgeDocumentRequest.content's ceiling, so extraction cannot
# produce something the store would then refuse.
MAX_EXTRACTED_CHARACTERS = 500_000

SUPPORTED_EXTENSIONS = frozenset(
    {".txt", ".md", ".csv", ".json", ".pdf", ".xlsx", ".xlsm", ".docx"}
)


class DocumentExtractionError(RuntimeError):
    """The file cannot be turned into text, with a reason a person can act on."""


def _clip(text: str) -> str:
    if len(text) <= MAX_EXTRACTED_CHARACTERS:
        return text
    return text[:MAX_EXTRACTED_CHARACTERS].rstrip() + "\n\n[…documento truncado…]"


def _extension(name: str) -> str:
    _, _, suffix = name.rpartition(".")
    return f".{suffix.lower()}" if suffix else ""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 maps every byte, so reaching here means something stranger.
    raise DocumentExtractionError(
        "No se pudo leer el texto del archivo: la codificación no es reconocible."
    )


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentExtractionError("Orion no tiene instalado el lector de PDF.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty password opens the common "protected but not secret"
            # case; anything else needs a password Orion does not have.
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentExtractionError(
                    "El PDF está protegido con contraseña y Orion no puede abrirlo."
                ) from exc
        pages = reader.pages[:MAX_PDF_PAGES]
        blocks = [f"[Página {index}]\n{(page.extract_text() or '').strip()}"
                  for index, page in enumerate(pages, start=1)]
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError(
            "El PDF no se pudo leer; puede estar dañado o tener un formato inusual."
        ) from exc

    text = "\n\n".join(block for block in blocks if block.split("\n", 1)[-1].strip())
    if not text.strip():
        # A scan is an image of text, not text. Saying so is more useful than
        # storing an empty document that will silently answer nothing.
        raise DocumentExtractionError(
            "El PDF no contiene texto seleccionable (suele pasar con documentos "
            "escaneados). Orion todavía no lee texto dentro de imágenes."
        )
    return _clip(text)


def _cell_text(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentExtractionError("Orion no tiene instalado el lector de Excel.") from exc

    try:
        # read_only streams rows instead of building the whole workbook, and
        # data_only takes the cached result of a formula rather than "=B2*1.5",
        # which is what a person reading the sheet actually sees.
        workbook = load_workbook(
            io.BytesIO(data), read_only=True, data_only=True, keep_links=False
        )
    except zipfile.BadZipFile as exc:
        raise DocumentExtractionError(
            "El archivo no es un Excel válido (.xlsx). Si es un .xls antiguo, "
            "guardalo como .xlsx o exportalo a CSV."
        ) from exc
    except Exception as exc:
        raise DocumentExtractionError(
            "El Excel no se pudo leer; puede estar dañado."
        ) from exc

    try:
        sections: list[str] = []
        for sheet in workbook.worksheets[:MAX_SHEETS]:
            buffer = io.StringIO()
            writer = csv.writer(buffer, lineterminator="\n")
            rows_written = 0
            for row in sheet.iter_rows(values_only=True):
                cells = [_cell_text(value) for value in row]
                if not any(cells):
                    continue
                writer.writerow(cells)
                rows_written += 1
                if rows_written >= MAX_SHEET_ROWS:
                    break
            if rows_written:
                sections.append((sheet.title, buffer.getvalue().rstrip()))
    finally:
        workbook.close()

    if not sections:
        raise DocumentExtractionError("El Excel no tiene filas con datos.")
    if len(sections) == 1:
        # One sheet becomes plain CSV, with no annotation of ours in it, so it
        # lands on the path Orion already computes over: filtering or averaging
        # a GPS export must not depend on whether the person exported it as
        # .csv or handed over the .xlsx they were sent.
        return _clip(sections[0][1])
    # Several sheets cannot be one table - their columns need not even match -
    # so they are labelled and kept as readable evidence. Orion can quote and
    # reason over them; the deterministic CSV tool works on a single-sheet
    # workbook or an exported sheet.
    return _clip(
        "\n\n".join(f"[Hoja: {title}]\n{body}" for title, body in sections)
    )


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentExtractionError("Orion no tiene instalado el lector de Word.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except (zipfile.BadZipFile, ValueError) as exc:
        raise DocumentExtractionError(
            "El archivo no es un Word válido (.docx). Si es un .doc antiguo, "
            "guardalo como .docx."
        ) from exc
    except Exception as exc:
        raise DocumentExtractionError(
            "El Word no se pudo leer; puede estar dañado."
        ) from exc

    blocks = [
        paragraph.text.strip()
        for paragraph in document.paragraphs[:MAX_DOCX_PARAGRAPHS]
        if paragraph.text.strip()
    ]
    # Tables carry the numbers in a report - a load summary, a test result - and
    # they live outside `paragraphs`, so reading only paragraphs would drop
    # exactly the part worth asking about.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    if not blocks:
        raise DocumentExtractionError("El documento de Word no contiene texto.")
    return _clip("\n\n".join(blocks))


def extract_text(filename: str, data: bytes) -> str:
    """Plain text from an uploaded file, or an error naming what went wrong.

    The format is chosen by the file's extension, which is the person's own
    declaration of what they uploaded - not by sniffing the bytes, and never by
    guessing from the content.
    """

    if not data:
        raise DocumentExtractionError("El archivo está vacío.")

    extension = _extension(filename)
    if extension in {".txt", ".md", ".csv", ".json"}:
        text = _decode_text(data)
        if not text.strip():
            raise DocumentExtractionError("El archivo está vacío.")
        return _clip(text)
    if extension == ".pdf":
        return _extract_pdf(data)
    if extension in {".xlsx", ".xlsm"}:
        return _extract_xlsx(data)
    if extension == ".docx":
        return _extract_docx(data)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise DocumentExtractionError(
        f"Orion todavía no puede leer archivos {extension or 'sin extensión'}. "
        f"Formatos aceptados: {supported}."
    )
